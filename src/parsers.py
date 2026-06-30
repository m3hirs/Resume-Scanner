import io
import pdfplumber
from docx import Document


def parse_pdf(file) -> str:
    """Extract text from a PDF file."""
    text_parts = []
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def parse_docx(file) -> str:
    """Extract text from a DOCX file."""
    doc = Document(file)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def parse_text(file) -> str:
    """Extract text from a plain text file or bytes."""
    if isinstance(file, bytes):
        return file.decode("utf-8", errors="replace")
    if isinstance(file, str):
        return file
    raw = file.read()
    if isinstance(raw, bytes):
        return raw.decode("utf-8", errors="replace")
    return raw


def parse_file(uploaded_file) -> str:
    """Dispatch to the correct parser based on file extension.

    Args:
        uploaded_file: A Streamlit UploadedFile object.

    Returns:
        Extracted text as a string.

    Raises:
        ValueError: If the file type is unsupported or extraction yields no text.
    """
    name = uploaded_file.name.lower()
    data = uploaded_file.read()
    buf = io.BytesIO(data)

    if name.endswith(".pdf"):
        text = parse_pdf(buf)
    elif name.endswith(".docx"):
        text = parse_docx(buf)
    elif name.endswith(".txt"):
        text = parse_text(data)
    else:
        raise ValueError(f"Unsupported file type: {name}. Please upload a PDF, DOCX, or TXT file.")

    if not text or not text.strip():
        raise ValueError("Could not extract any text from the uploaded file.")

    return text
